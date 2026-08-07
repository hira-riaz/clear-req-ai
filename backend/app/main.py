"""
ClearReq AI backend.

Endpoints map onto the pipeline diagrams plus the full session workflow:

  Start session:      POST /sessions
  Detection phase:     POST /requirements/analyze
  Resolution phase:    POST /requirements/translate
  Review/edit:          PATCH /requirements/{id}/edit
  Report (data):        GET  /sessions/{id}/report
  Report (Word file):    GET  /sessions/{id}/report/docx
"""
import io
from datetime import datetime

from fastapi import FastAPI, Depends
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session as DBSession
from docx import Document

from . import models, rule_detector, ai_provider
from .database import engine, get_db
from .schemas import RequirementIn, TranslateRequest, SessionIn, RequirementEdit

models.Base.metadata.create_all(bind=engine)

app = FastAPI(title="ClearReq AI")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],  # tighten this before any real deployment
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/")
def root():
    return {"status": "ClearReq AI backend running"}


@app.post("/sessions")
def start_session(payload: SessionIn, db: DBSession = Depends(get_db)):
    """Workflow step 1: create a project + session to group requirements."""
    project = models.Project(name=payload.project_name, client_name=payload.client_name)
    db.add(project)
    db.commit()
    db.refresh(project)

    session = models.Session(project_id=project.id)
    db.add(session)
    db.commit()
    db.refresh(session)

    return {"session_id": session.id, "project_id": project.id, "project_name": project.name}


def _merge_ambiguities(rule_results: list[dict], ai_results: list[dict]) -> list[dict]:
    """Deduplicate by term, preferring the rule-based entry when both agree
    (it's deterministic and free), otherwise keeping each unique catch."""
    by_term = {r["term"].lower(): r for r in rule_results}
    for r in ai_results:
        key = r["term"].lower()
        if key not in by_term:
            by_term[key] = r
    return list(by_term.values())


def _find_previous_answer(db: DBSession, session_id: int, exclude_requirement_id: int, term: str) -> str | None:
    """
    Session memory: if this exact term was already clarified earlier in the
    same session, return that answer so the user isn't asked to re-answer
    something they've already resolved. Frontend pre-fills it, editable.
    """
    result = (
        db.query(models.Clarification)
        .join(models.Ambiguity, models.Clarification.ambiguity_id == models.Ambiguity.id)
        .join(models.Requirement, models.Ambiguity.requirement_id == models.Requirement.id)
        .filter(models.Requirement.session_id == session_id)
        .filter(models.Requirement.id != exclude_requirement_id)
        .filter(models.Ambiguity.term.ilike(term))
        .filter(models.Clarification.answer.isnot(None))
        .order_by(models.Clarification.answered_at.desc())
        .first()
    )
    return result.answer if result else None


@app.post("/requirements/analyze")
def analyze_requirement(payload: RequirementIn, db: DBSession = Depends(get_db)):
    """
    Detection phase: run both detectors, merge/deduplicate, check for
    conflicts with other requirements in the session, check for terms
    already clarified earlier in the session (session memory), save
    everything, and return it to the frontend.
    """
    requirement = models.Requirement(
        session_id=payload.session_id,
        original_text=payload.text,
        status="clarifying",
    )
    db.add(requirement)
    db.commit()
    db.refresh(requirement)

    rule_results = rule_detector.detect(payload.text)
    ai_results = ai_provider.detect_ambiguity(payload.text)
    merged = _merge_ambiguities(rule_results, ai_results)

    # Session-level consistency: check for conflicts with already-translated requirements
    existing_versions = (
        db.query(models.RequirementVersion)
        .join(models.Requirement)
        .filter(models.Requirement.session_id == payload.session_id)
        .filter(models.Requirement.id != requirement.id)
        .all()
    )
    existing_texts = [v.translated_text for v in existing_versions]
    conflicts = ai_provider.check_conflicts(payload.text, existing_texts)
    for c in conflicts:
        merged.append({
            "term": f"conflict with: {c['conflicts_with'][:60]}",
            "category": "conflict",
            "detector": "ai",
            "confidence": 1.0,
            "question": c["question"],
        })

    saved = []
    for item in merged:
        ambiguity = models.Ambiguity(
            requirement_id=requirement.id,
            term=item["term"],
            category=item["category"],
            detector=item["detector"],
            confidence=item["confidence"],
        )
        db.add(ambiguity)
        db.commit()
        db.refresh(ambiguity)

        clarification = models.Clarification(
            ambiguity_id=ambiguity.id,
            question=item["question"],
        )
        db.add(clarification)
        db.commit()

        # Session memory: was this exact term already clarified earlier?
        suggested_answer = None
        if item["category"] != "conflict":
            suggested_answer = _find_previous_answer(db, payload.session_id, requirement.id, item["term"])

        saved.append({
            "ambiguity_id": ambiguity.id,
            "term": ambiguity.term,
            "category": ambiguity.category,
            "detector": ambiguity.detector,
            "confidence": ambiguity.confidence,
            "question": clarification.question,
            "suggested_answer": suggested_answer,
        })

    return {"requirement_id": requirement.id, "ambiguities": saved}


@app.post("/requirements/translate")
def translate_requirement(payload: TranslateRequest, db: DBSession = Depends(get_db)):
    """
    Resolution phase: save clarification answers, call the AI provider to
    compose the final translation (using prior session requirements as
    context for consistency), score confidence, save a new version.
    """
    requirement = db.get(models.Requirement, payload.requirement_id)

    clarifications_for_prompt = []
    for ans in payload.answers:
        clarification = (
            db.query(models.Clarification)
            .filter(models.Clarification.ambiguity_id == ans.ambiguity_id)
            .first()
        )
        clarification.answer = ans.answer
        clarification.answered_at = datetime.utcnow()
        db.commit()
        clarifications_for_prompt.append({
            "term": clarification.ambiguity.term,
            "question": clarification.question,
            "answer": ans.answer,
        })

    context_versions = (
        db.query(models.RequirementVersion)
        .join(models.Requirement)
        .filter(models.Requirement.session_id == requirement.session_id)
        .filter(models.Requirement.id != requirement.id)
        .all()
    )
    context_texts = [v.translated_text for v in context_versions]

    result = ai_provider.translate(requirement.original_text, clarifications_for_prompt, context_texts)

    existing_versions = (
        db.query(models.RequirementVersion)
        .filter(models.RequirementVersion.requirement_id == requirement.id)
        .count()
    )
    version = models.RequirementVersion(
        requirement_id=requirement.id,
        version_number=existing_versions + 1,
        translated_text=result["translated_text"],
        confidence_score=result["confidence"],
    )
    db.add(version)
    requirement.status = "translated"
    db.commit()
    db.refresh(version)

    return {
        "requirement_id": requirement.id,
        "version_number": version.version_number,
        "translated_text": version.translated_text,
        "confidence_score": version.confidence_score,
    }


@app.patch("/requirements/{requirement_id}/edit")
def edit_requirement_translation(requirement_id: int, payload: RequirementEdit, db: DBSession = Depends(get_db)):
    """
    Lets the user hand-edit a translated requirement during review, before
    the final report is generated. Creates a new version rather than
    overwriting, so edit history is preserved.
    """
    existing_versions = (
        db.query(models.RequirementVersion)
        .filter(models.RequirementVersion.requirement_id == requirement_id)
        .count()
    )
    version = models.RequirementVersion(
        requirement_id=requirement_id,
        version_number=existing_versions + 1,
        translated_text=payload.translated_text,
        confidence_score=1.0,  # human-edited, treated as fully confident
    )
    db.add(version)
    db.commit()
    db.refresh(version)
    return {
        "requirement_id": requirement_id,
        "version_number": version.version_number,
        "translated_text": version.translated_text,
    }


@app.get("/requirements/{requirement_id}")
def get_requirement(requirement_id: int, db: DBSession = Depends(get_db)):
    requirement = db.get(models.Requirement, requirement_id)
    versions = (
        db.query(models.RequirementVersion)
        .filter(models.RequirementVersion.requirement_id == requirement_id)
        .all()
    )
    return {
        "id": requirement.id,
        "original_text": requirement.original_text,
        "status": requirement.status,
        "versions": [
            {"version_number": v.version_number, "translated_text": v.translated_text,
             "confidence_score": v.confidence_score}
            for v in versions
        ],
    }


@app.get("/sessions/{session_id}/report")
def get_session_report(session_id: int, db: DBSession = Depends(get_db)):
    """Every requirement in the session with its latest translated version."""
    session = db.get(models.Session, session_id)
    project = db.get(models.Project, session.project_id) if session else None

    requirements = (
        db.query(models.Requirement)
        .filter(models.Requirement.session_id == session_id)
        .all()
    )

    items = []
    for r in requirements:
        latest = (
            db.query(models.RequirementVersion)
            .filter(models.RequirementVersion.requirement_id == r.id)
            .order_by(models.RequirementVersion.version_number.desc())
            .first()
        )
        items.append({
            "requirement_id": r.id,
            "original_text": r.original_text,
            "status": r.status,
            "translated_text": latest.translated_text if latest else None,
            "confidence_score": latest.confidence_score if latest else None,
        })

    return {
        "session_id": session_id,
        "project_name": project.name if project else None,
        "requirements": items,
    }


@app.get("/sessions/{session_id}/report/docx")
def download_report_docx(session_id: int, db: DBSession = Depends(get_db)):
    """
    Same report as /report above, rendered as a downloadable Word document:
    numbered translated requirements, followed by an appendix with the
    client's original wording preserved exactly.
    """
    session = db.get(models.Session, session_id)
    project = db.get(models.Project, session.project_id) if session else None
    requirements = (
        db.query(models.Requirement)
        .filter(models.Requirement.session_id == session_id)
        .all()
    )

    doc = Document()
    title = project.name if project else "ClearReq AI Report"
    doc.add_heading(f"{title} — System Requirements Specification", level=1)
    doc.add_paragraph(f"Generated by ClearReq AI on {datetime.utcnow().strftime('%Y-%m-%d')}")

    doc.add_heading("Translated Requirements", level=2)
    for r in requirements:
        latest = (
            db.query(models.RequirementVersion)
            .filter(models.RequirementVersion.requirement_id == r.id)
            .order_by(models.RequirementVersion.version_number.desc())
            .first()
        )
        text = latest.translated_text if latest else "(no translation)"
        doc.add_paragraph(text, style="List Number")

    doc.add_page_break()
    doc.add_heading("Appendix: Original Client Requirements", level=2)
    for r in requirements:
        doc.add_paragraph(r.original_text, style="List Number")

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    safe_title = (project.name if project else "clearreq").replace(" ", "_")
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={safe_title}_report.docx"},
    )